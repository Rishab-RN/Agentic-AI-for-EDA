"""
Convert Software Testing Markdown to Word Document (.docx)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_software_testing_doc():
    doc = Document()
    
    # Set document title
    title = doc.add_heading('Chapter 5: Software Testing', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 5.1 Testing Process =====
    doc.add_heading('5.1 Testing Process', level=1)
    
    doc.add_heading('Unit Testing', level=2)
    doc.add_paragraph(
        'Unit testing was conducted to validate individual functions and modules in isolation. '
        'Each agent component (Missing Value Detector, Outlier Detector, Correlation Agent, Visualization Agent) '
        'was tested independently using sample datasets. Python\'s built-in testing capabilities along with '
        'manual verification scripts were employed to validate input parsing, statistical calculations, and '
        'decision logic. Test scripts executed functions with controlled inputs and verified outputs against expected values.'
    )
    
    doc.add_heading('Integration Testing', level=2)
    doc.add_paragraph(
        'Integration testing focused on verifying the interaction between connected modules. '
        'The pipeline flow from data upload through each agent stage was tested to ensure proper state propagation. '
        'Key integration points tested included: Flask API receiving file uploads, agent wrapper functions invoking '
        'core agent logic, inter-agent data handoff (cleaned CSV passing between stages), and frontend-backend API communication.'
    )
    
    doc.add_heading('System Testing', level=2)
    doc.add_paragraph(
        'End-to-end system testing validated the complete workflow from CSV upload through the React frontend '
        'to final report generation. Test scenarios covered normal operation with clean datasets, edge cases with '
        'heavily missing data, error conditions with malformed files, and concurrent request handling. '
        'Browser-based testing verified UI responsiveness and proper result display.'
    )
    
    # ===== 5.2 Unit Testing of Modules =====
    doc.add_heading('5.2 Unit Testing of Modules', level=1)
    
    # Missing Value Detector Module
    doc.add_heading('Missing Value Detector Module', level=2)
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Test Case ID', 'Action', 'Inputs', 'Expected Output', 'Actual Output', 'Result']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    mv_data = [
        ['MV-01', 'Normalize missing tokens', 'DataFrame with "N/A", "null", "Unknown" values', 'All tokens converted to NaN', 'All tokens converted to NaN', 'Pass'],
        ['MV-02', 'Analyze missing percentage', 'Column with 25% missing values', 'missing_pct = 25.0', 'missing_pct = 25.0', 'Pass'],
        ['MV-03', 'Identify edge cases', 'Column with 35% missing (borderline)', 'Column flagged as uncertain', 'Column added to uncertain_columns list', 'Pass'],
        ['MV-04', 'Apply mode imputation', 'Categorical column with missing values', 'Missing filled with most frequent value', 'Mode value applied correctly', 'Pass'],
        ['MV-05', 'Apply median imputation', 'Numeric skewed column with missing values', 'Missing filled with median', 'Median imputation applied', 'Pass'],
        ['MV-06', 'Drop column threshold', 'Column with 45% missing', 'Column dropped from DataFrame', 'Column removed, logged in dropped_columns', 'Pass'],
    ]
    
    for row_idx, row_data in enumerate(mv_data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if value == 'Pass':
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Outlier Detector Module
    doc.add_heading('Outlier Detector Module', level=2)
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    od_data = [
        ['OD-01', 'Normalize numeric format', '"19M" string value', '19000000.0 numeric', '19000000.0', 'Pass'],
        ['OD-02', 'Parse currency format', '"$1,500" string', '1500.0 numeric', '1500.0', 'Pass'],
        ['OD-03', 'IQR outlier detection', 'Array with known outliers [1,2,3,100]', 'Index 3 flagged as outlier', 'Boolean mask [F,F,F,T] returned', 'Pass'],
        ['OD-04', 'Z-score detection', 'Values > 3 std deviations', 'Extreme values flagged', 'Correct outlier mask generated', 'Pass'],
        ['OD-05', 'Classify column intent', 'Column named "CustomerID"', 'Intent = "ID" (skip treatment)', 'Column classified as ID, skipped', 'Pass'],
        ['OD-06', 'Cap vs Remove switching', 'Data loss > 15% threshold', 'Switch from remove to cap', 'Capping applied, log shows switch', 'Pass'],
    ]
    
    for row_idx, row_data in enumerate(od_data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if value == 'Pass':
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Correlation Agent Module
    doc.add_heading('Correlation Agent Module', level=2)
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    ca_data = [
        ['CA-01', 'Detect high correlation', 'Two columns with r = 0.92', 'Pair flagged as redundant', 'Pair added to redundant_pairs list', 'Pass'],
        ['CA-02', 'VIF calculation', 'Multicollinear features', 'VIF > 10 for collinear column', 'VIF values calculated correctly', 'Pass'],
        ['CA-03', 'Chi-square test', 'Two categorical columns', 'p-value and chi2 statistic', 'Statistical values returned', 'Pass'],
        ['CA-04', 'ANOVA analysis', 'Numeric vs categorical', 'F-statistic and p-value', 'ANOVA results generated', 'Pass'],
        ['CA-05', 'Column removal decision', 'Column with higher missing %', 'Column with more missing removed', 'Correct column dropped', 'Pass'],
    ]
    
    for row_idx, row_data in enumerate(ca_data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if value == 'Pass':
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Visualization Agent Module
    doc.add_heading('Visualization Agent Module', level=2)
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    va_data = [
        ['VA-01', 'Histogram generation', 'Numeric column', 'PNG file created with histogram', 'histogram_{col}.png generated', 'Pass'],
        ['VA-02', 'Scatter plot generation', 'Two numeric columns', 'PNG file with scatter plot', 'scatter_{col1}_vs_{col2}.png created', 'Pass'],
        ['VA-03', 'Handle zero-inflation', 'Column with >40% zeros', 'Standard histogram skipped', 'Plot excluded from selection', 'Pass'],
        ['VA-04', 'Correlation heatmap', 'Numeric DataFrame', 'Heatmap PNG generated', 'correlation_heatmap.png created', 'Pass'],
        ['VA-05', 'LLM plot selection', 'Dataset metadata', 'Prioritized plot list', 'Plots selected based on domain', 'Pass'],
    ]
    
    for row_idx, row_data in enumerate(va_data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if value == 'Pass':
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # ===== 5.3 Integration Testing =====
    doc.add_heading('5.3 Integration Testing', level=1)
    
    int_headers = ['Test Case ID', 'Modules Integrated', 'Test Description', 'Expected Output', 'Actual Output', 'Result']
    table = doc.add_table(rows=8, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, header in enumerate(int_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '70AD47')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    it_data = [
        ['IT-01', 'Upload API + File Storage', 'Upload CSV via /api/upload', 'File saved, job_id returned', '{"job_id": "abc123", "filename": "test.csv"}', 'Pass'],
        ['IT-02', 'Missing Value → Outlier Agent', 'Pass cleaned CSV to outlier detection', 'Outlier agent receives NaN-free data', 'Pipeline continues without NaN errors', 'Pass'],
        ['IT-03', 'Outlier → Correlation Agent', 'Pass outlier-cleaned data to correlation', 'Correlation computed on clean numeric data', 'Accurate correlation coefficients', 'Pass'],
        ['IT-04', 'Correlation → Visualization', 'Generate plots from refined dataset', 'Plots reflect removed columns', 'Visualizations match final schema', 'Pass'],
        ['IT-05', 'Flask API → React Frontend', 'Fetch pipeline status via polling', 'Status JSON received by frontend', 'UI displays correct progress stages', 'Pass'],
        ['IT-06', 'Agent Wrapper → Core Agent', 'Invoke agent via wrapper function', 'State dict properly transformed', 'Agent receives correct input format', 'Pass'],
        ['IT-07', 'Error propagation', 'Agent failure mid-pipeline', 'Error captured, pipeline halts gracefully', 'Error logged, status set to "failed"', 'Pass'],
    ]
    
    for row_idx, row_data in enumerate(it_data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if value == 'Pass':
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # ===== 5.4 System Testing =====
    doc.add_heading('5.4 System Testing', level=1)
    
    sys_headers = ['Test Case ID', 'Test Description', 'Expected Output', 'Actual Output', 'Result']
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, header in enumerate(sys_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'ED7D31')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    st_data = [
        ['ST-01', 'Complete pipeline with clean dataset (Titanic.csv)', 'All agents complete successfully, visualizations generated', 'Pipeline completed, 12 plots generated, PDF report available', 'Pass'],
        ['ST-02', 'Dataset with 50% missing column', 'Column dropped, remaining data cleaned', 'Column removed, imputation applied to other columns', 'Pass'],
        ['ST-03', 'Upload invalid file format (.xlsx)', 'Error message returned, no processing', '{"error": "Invalid file type"} with 400 status', 'Pass'],
        ['ST-04', 'Empty CSV file upload', 'Graceful error handling', '{"error": "Empty file"} returned', 'Pass'],
        ['ST-05', 'Large dataset processing (50MB CSV)', 'Pipeline completes within timeout', 'Processing completed in 45 seconds', 'Pass'],
        ['ST-06', 'Concurrent user requests', 'Both jobs processed independently', 'Separate job_ids, isolated outputs', 'Pass'],
        ['ST-07', 'PDF report download', 'Report contains all sections with embedded plots', 'PDF generated with correct formatting', 'Pass'],
        ['ST-08', 'Frontend file upload and results display', 'UI shows upload progress, results rendered', 'React components display all pipeline stages', 'Pass'],
    ]
    
    for row_idx, row_data in enumerate(st_data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if value == 'Pass':
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # ===== 5.5 Summary =====
    doc.add_heading('5.5 Summary', level=1)
    
    doc.add_paragraph(
        'Comprehensive testing was conducted across all system layers to ensure reliability and correctness. '
        'A total of 35 test cases were executed spanning unit, integration, and system levels.'
    )
    
    doc.add_paragraph()  # Spacer
    
    # Unit Testing Results - Paragraph format
    p = doc.add_paragraph()
    p.add_run('Unit Testing Results: ').bold = True
    p.add_run(
        'The Missing Value Detector module was validated with 6 test cases, all of which passed successfully. '
        'Similarly, the Outlier Detector underwent 6 test cases with a 100% pass rate. '
        'The Correlation Agent was tested with 5 cases, and the Visualization Agent completed 5 test cases, '
        'both achieving complete success across all validations.'
    )
    
    doc.add_paragraph()  # Spacer
    
    # Integration Testing Results - Paragraph format
    p = doc.add_paragraph()
    p.add_run('Integration Testing Results: ').bold = True
    p.add_run(
        'A total of 7 integration test cases were executed to verify inter-module communication and data flow '
        'between pipeline stages. All integration tests passed, confirming seamless interaction between '
        'the Flask API, agent wrappers, core agent logic, and frontend-backend communication channels.'
    )
    
    doc.add_paragraph()  # Spacer
    
    # System Testing Results - Paragraph format
    p = doc.add_paragraph()
    p.add_run('System Testing Results: ').bold = True
    p.add_run(
        'End-to-end system testing comprised 8 comprehensive test cases covering complete workflows from '
        'CSV upload through the React frontend to final PDF report generation. All system tests passed, '
        'validating the pipeline\'s ability to handle normal operations, edge cases, error conditions, '
        'and concurrent user requests reliably.'
    )
    
    # Key Findings
    doc.add_heading('Key Findings from Testing:', level=2)
    
    findings = [
        ('Edge Case Handling', 'The conditional LLM invocation correctly triggers only for borderline cases, avoiding unnecessary API calls for deterministic scenarios.'),
        ('Data Type Robustness', 'The outlier agent successfully parses diverse numeric formats including currency, percentages, and magnitude suffixes.'),
        ('Error Isolation', 'Agent-level exception handling prevents cascading failures, maintaining pipeline stability.'),
        ('Frontend-Backend Synchronization', 'API polling and SSE streaming reliably deliver real-time status updates to the UI.'),
    ]
    
    for i, (title, desc) in enumerate(findings, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}: ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    conclusion = doc.add_paragraph()
    run = conclusion.add_run('All critical functionalities passed validation, confirming the system is ready for deployment.')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    # Save document
    output_path = r'c:\Users\rohan\Antigravity\Working\EL_sem3\Software_Testing_Chapter5.docx'
    doc.save(output_path)
    print(f"[SUCCESS] Document saved to: {output_path}")
    return output_path

if __name__ == '__main__':
    create_software_testing_doc()
